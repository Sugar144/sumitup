#!/usr/bin/env python3
"""
Generate one DOCX per chapter (Fundamentals of SA, Part I, Chapters 1–8).

Source: pdftotext -layout on the publisher PDF in SA/.
Output: data/sa/fundamentals/extraction_docx/Ch##_*.docx

Each document mirrors the book section order. For every section:
  - Heading (book title)
  - Table: Subpoint | Extracted text (from PDF) | Study note (plain-language gloss)
  - Full section body as appendix paragraphs where useful

Requires: python-docx, pdftotext (poppler)
"""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
PDF = ROOT / "SA/dokumen.pub_fundamentals-of-software-architecture-a-modern-engineering-approach-2.pdf"
OUT_DIR = ROOT / "data/sa/fundamentals/extraction_docx"
TMP_TXT = ROOT / ".fsa_extract.txt"

# Line ranges are 1-based inclusive start, inclusive end, from pdftotext -layout output
CHAPTER_RANGES: list[tuple[int, int, int, str]] = [
    (1, 802, 1211, "Introduction"),
    (2, 1212, 1878, "Architectural Thinking"),
    (3, 1879, 2533, "Modularity"),
    (4, 2534, 2964, "Architectural Characteristics Defined"),
    (5, 2965, 3442, "Identifying Architectural Characteristics"),
    (6, 3443, 3927, "Measuring and Governing Architecture Characteristics"),
    (7, 3928, 4297, "The Scope of Architectural Characteristics"),
    (8, 4298, 5004, "Component-Based Thinking"),
]

# Ordered section markers exactly as they appear in pdftotext (trimmed lines)
SECTION_MARKERS: dict[int, list[str]] = {
    1: [
        "Defining Software Architecture",
        "Laws of Software Architecture",
        "Expectations of an Architect",
        "Make Architecture Decisions",
        "Continually Analyze the Architecture",
        "Keep Current with Latest Trends",
        "Ensure Compliance with Decisions",
        "Understand Diverse Technologies",
        "Know the Business Domain",
        "Possess Interpersonal Skills",
        "Understand and Navigate Politics",
        "Roadmap",
    ],
    2: [
        "Architecture Versus Design",
        "Strategic Versus Tactical Decisions",
        "Level of Effort",
        "The Significance of Trade-Offs",
        "Technical Breadth",
        "The 20-Minute Rule",
        "Developing a Personal Radar",
        "Analyzing Trade-Offs",
        "Understanding Business Drivers",
        "Balancing Architecture and Hands-On Coding",
        "There’s More to Architectural Thinking",
    ],
    3: [
        "Modularity Versus Granularity",
        "Defining Modularity",
        "Measuring Modularity",
        "Cohesion",
        "Coupling",
        "Core Metrics",
        "Distance from the Main Sequence",
        "Connascence",
        "From Modules to Components",
    ],
    4: [
        "The Longevity of the Term “Non-Functional Requirements”",
        "Architectural Characteristics and System Design",
        "Architectural Characteristics (Partially) Listed",
        "Operational Architectural Characteristics",
        "Structural Architectural Characteristics",
        "Cloud Characteristics",
        "Cross-Cutting Architectural Characteristics",
        "The Many Ambiguities in Software Architecture",
        "Trade-Offs and Least Worst Architecture",
    ],
    5: [
        "from Domain Concerns",
        "Composite Architectural Characteristics",
        "The Origin of Architecture Katas",
        "Working with Katas",
        "Kata: Silicon Sandwiches",
        "Explicit Characteristics",
        "Implicit Characteristics",
        "Design Versus Architecture and Trade-Offs",
        "Limiting and Prioritizing Architectural Characteristics",
        "Case Study: The Vasa",
    ],
    6: [
        "Measuring Architecture Characteristics",
        "Operational Measures",
        "The Many Flavors of Performance",
        "Structural Measures",
        "Cyclomatic Complexity",
        "Process Measures",
        "Governance and Fitness Functions",
        "Governing Architecture Characteristics",
        "Fitness Functions",
    ],
    7: [
        "Architectural Quanta and Granularity",
        "Domain-Driven Design’s Bounded Context",
        "Synchronous Communication",
        "The Impact of Scoping",
        "Scoping and Architectural Style",
        "Kata: Going Green",
        "Scoping and the Cloud",
    ],
    8: [
        "Defining Logical Components",
        "Logical Versus Physical Architecture",
        "Creating a Logical Architecture",
        "Identifying Core Components",
        "Assigning User Stories to Components",
        "Analyzing Roles and Responsibilities",
        "Analyzing Architectural Characteristics",
        "Restructuring Components",
        "Component Coupling",
        "Static Coupling",
        "Temporal Coupling",
        "The Law of Demeter",
        "Case Study: Going, Going, Gone—Discovering Components",
    ],
}

# Short study gloss per chapter/section key (section title -> note). Fills third column of tables.
STUDY_NOTES: dict[int, dict[str, str]] = {
    1: {
        "Defining Software Architecture": "Four dimensions work together: what the system must do beyond features (-ilities), how behaviour is grouped (components), which style you start from, and written rules (decisions). Architects analyse characteristics and components before picking a style.",
        "Laws of Software Architecture": "If it looks like a free lunch, you missed a trade-off. Document why, not only how. Decisions are usually gradients, not switches.",
        "Expectations of an Architect": "The job mixes technical leadership with organisational work: set guard rails, verify teams follow them, stay broad technically, and speak both business and engineering languages.",
        "Make Architecture Decisions": "Prefer principles that guide (e.g. “use a reactive UI framework”) over picking one vendor stack, unless a specific tech is required to preserve an -ility.",
        "Continually Analyze the Architecture": "Watch for structural decay and for “ agility in code” killed by slow test/release pipelines.",
        "Keep Current with Latest Trends": "Architect decisions last years; understanding cloud, AI, and platform shifts reduces future regret.",
        "Ensure Compliance with Decisions": "Without checks, shortcuts (e.g. Presentation→DB) undermine the very characteristics you designed for.",
        "Understand Diverse Technologies": "Breadth beats being the world’s top expert in one cache product; know several options and their trade-offs.",
        "Know the Business Domain": "You cannot design credibly for trading, health, or logistics without understanding how stakeholders talk and count risk.",
        "Possess Interpersonal Skills": "Architecture is enacted through teams; facilitation and leadership are as important as diagrams.",
        "Understand and Navigate Politics": "Big structural moves affect budgets and other teams; expect negotiation, unlike local refactor choices.",
        "Roadmap": "Part I builds analysis skills; Part II catalogues styles; Part III covers decisions, risk, diagrams, and soft skills.",
    },
    2: {
        "Architecture Versus Design": "Structure vs look-and-feel; many real choices sit in between—use strategy, cost to change, and size of trade-offs to judge.",
        "Strategic Versus Tactical Decisions": "Planning depth, number of stakeholders, and time horizon hint whether you are shaping the system or detailing it.",
        "Level of Effort": "Fowler: architecture is the hard-to-change stuff; cheap reversals usually sit on the design side.",
        "The Significance of Trade-Offs": "Microservices vs monolith is architectural because the trade-offs are huge; splitting a class file is design-level.",
        "Technical Breadth": "Move knowledge from unknown-unknowns to known-unknowns, then deepen only where needed.",
        "The 20-Minute Rule": "Daily, before email, invest a small block to scan radars and terms—compound learning.",
        "Developing a Personal Radar": "Hold/Assess/Trial/Adopt rings structure where you place tools and practices; diversifies your career “portfolio”.",
        "Analyzing Trade-Offs": "“It depends” is correct; context (org, budget, skills, runtime) picks queues vs topics, sync vs async, etc.",
        "Understanding Business Drivers": "Time to market, compliance, cost, growth—each stresses different -ilities.",
        "Balancing Architecture and Hands-On Coding": "Stay credible with PoCs, debt fixes, tooling, and reviews without becoming the bottleneck implementer.",
        "There’s More to Architectural Thinking": "Chapter closes the loop: thinking architecturally is continuous, not a one-off diagram.",
    },
    # ... filled below via merge
}

# Merge minimal defaults for sections without explicit note
_DEFAULT = "Relate this subsection to trade-offs: what changes if you neglect this idea?"

# When PDF splits a heading across lines or uses a short anchor, use this for the DOCX heading text.
SECTION_DISPLAY_TITLES: dict[int, dict[str, str]] = {
    5: {
        "from Domain Concerns": "Extracting Architectural Characteristics from Domain Concerns",
    },
}


def norm_line(s: str) -> str:
    s = s.strip()
    s = s.replace("\u2019", "'").replace("\u2018", "'")
    s = s.replace("\u201c", '"').replace("\u201d", '"')
    s = s.replace("\u2014", "-").replace("\u2013", "-")
    return s


def ensure_extract() -> list[str]:
    if not PDF.exists():
        raise SystemExit(f"Missing PDF: {PDF}")
    subprocess.run(
        ["pdftotext", "-layout", str(PDF), str(TMP_TXT)],
        check=True,
    )
    return TMP_TXT.read_text(encoding="utf-8", errors="replace").splitlines()


def clean_snippet(text: str, max_len: int = 3500) -> str:
    text = re.sub(r"\f", "\n", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" +\n", "\n", text)
    text = text.strip()
    if len(text) > max_len:
        text = text[: max_len - 3] + "..."
    return text


def find_marker_line(lines: list[str], marker: str, start: int = 0) -> int:
    m = norm_line(marker)
    if not m:
        return -1
    for i in range(start, len(lines)):
        nl = norm_line(lines[i])
        if nl == m:
            return i
        # Long headings may be truncated in layout export; short titles need exact match only.
        if len(m) >= 12 and nl.startswith(m):
            return i
    return -1


def split_chapter(
    lines: list[str], markers: list[str], ch: int
) -> list[tuple[str, str]]:
    """Return [(section_title, body_text), ...]. First chunk before first marker is opening."""
    chapter = "\n".join(lines)
    display = SECTION_DISPLAY_TITLES.get(ch, {})

    sections: list[tuple[str, str]] = []
    indices: list[tuple[int, str]] = []
    L = lines
    pos = 0
    for mk in markers:
        idx = find_marker_line(L, mk, pos)
        if idx >= 0:
            indices.append((idx, mk))
            pos = idx + 1

    if not indices:
        sections.append(("Chapter opening", clean_snippet(chapter)))
        return sections

    # Opening before first marker
    first_i = indices[0][0]
    opening = "\n".join(L[:first_i])
    if opening.strip():
        sections.append(("Chapter opening", clean_snippet(opening)))

    for j, (idx, mk) in enumerate(indices):
        end = indices[j + 1][0] if j + 1 < len(indices) else len(L)
        body = "\n".join(L[idx:end])
        title = display.get(mk, mk)
        sections.append((title, clean_snippet(body)))

    return sections


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, val in enumerate(row):
            cells[c].text = val


def build_chapter_doc(ch: int, start: int, end: int, title: str, all_lines: list[str]):
    lines = all_lines[start - 1 : end]
    markers = SECTION_MARKERS.get(ch, [])
    sections = split_chapter(lines, markers, ch)
    notes = STUDY_NOTES.get(ch, {})
    display = SECTION_DISPLAY_TITLES.get(ch, {})

    def note_for(sec_title: str) -> str:
        if sec_title in notes:
            return notes[sec_title]
        ns = norm_line(sec_title)
        for nk, v in notes.items():
            if norm_line(nk) == ns:
                return v
        for mk, disp in display.items():
            if norm_line(disp) == ns or norm_line(mk) == ns:
                if mk in notes:
                    return notes[mk]
                for nk, v in notes.items():
                    if norm_line(nk) == norm_line(mk):
                        return v
        return _DEFAULT

    doc = Document()
    doc.add_heading(f"Chapter {ch}: {title}", 0)
    p = doc.add_paragraph(
        "Fundamentals of Software Architecture (2e) — Richards & Ford. "
        "Knowledge extraction for personal study. Text excerpts from pdftotext layout export."
    )
    for run in p.runs:
        run.font.size = Pt(10)

    doc.add_paragraph(
        "Structure follows the book section order. Column “Study note” is a concise reading aid, not a substitute for the full chapter."
    )

    for sec_title, body in sections:
        doc.add_heading(sec_title, level=1)
        note = note_for(sec_title)
        # Table: one row with extract + note (subpoints stay inside extract)
        add_table(
            doc,
            ["Section", "Extracted text (from PDF)", "Study note (plain language)"],
            [[sec_title[:80], body[:8000] if body else "(no body)", note[:2000]]],
        )
        if len(body) > 8000:
            doc.add_heading("Full extract (continued)", level=2)
            doc.add_paragraph(body[8000:16000] if len(body) > 8000 else "")

    doc.add_page_break()
    doc.add_heading("Annex — full chapter plain text", level=1)
    full = clean_snippet("\n".join(lines), max_len=120_000)
    for chunk in range(0, len(full), 9000):
        doc.add_paragraph(full[chunk : chunk + 9000])

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    fname = OUT_DIR / f"Ch{ch:02d}_{title.replace(' ', '_').replace('/', '-')}_Knowledge_Extraction.docx"
    doc.save(fname)
    print("Wrote", fname)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    lines = ensure_extract()

    # Ch8: Case Study title (em dash in PDF)
    SECTION_MARKERS[8][-1] = "Case Study: Going, Going, Gone—Discovering Components"

    # Expand STUDY_NOTES for chapters 3-8 (compact)
    STUDY_NOTES[3] = {
        "Modularity Versus Granularity": "Modularity = splitting into parts; granularity = size of each part. Too-fine services → distributed mud.",
        "Defining Modularity": "Logical grouping of related code (packages/namespaces); not always physical deploy units.",
        "Measuring Modularity": "Use cohesion, coupling, and derived metrics rather than gut feel only.",
        "Cohesion": "How strongly parts of a module belong together; splitting cohesive modules hurts.",
        "Coupling": "Afferent (fan-in) vs efferent (fan-out); both matter for change risk.",
        "Core Metrics": "Abstractness, Instability, Distance from Main Sequence summarise maintainability pressure.",
        "Distance from the Main Sequence": "Balance abstract vs volatile; zones of pain/uselessness flag refactor targets.",
        "Connascence": "Vocabulary of dependency kinds; prefer weaker forms and tighter locality.",
        "From Modules to Components": "Components are the architectural name for deployable/reusable bundles built from modules.",
    }
    STUDY_NOTES[4] = {
        "The Longevity of the Term “Non-Functional Requirements”": "Authors prefer architectural characteristics / capabilities over belittling “non-functional” language.",
        "Architectural Characteristics and System Design": "Three tests: non-domain, structural impact, critical to success; explicit vs implicit.",
        "Architectural Characteristics (Partially) Listed": "Lists are exemplars—each org should define its ubiquitous language.",
        "Operational Architectural Characteristics": "Runtime qualities: scale, performance, availability, recovery…",
        "Structural Architectural Characteristics": "Codebase qualities: maintainability, modularity, portability…",
        "Cloud Characteristics": "Elasticity, regions, zones—provider capabilities become design constraints.",
        "Cross-Cutting Architectural Characteristics": "Security, legal, privacy span layers; often need explicit architecture support.",
        "The Many Ambiguities in Software Architecture": "Terms overlap; align with stakeholders and document definitions.",
        "Trade-Offs and Least Worst Architecture": "Optimising everything is impossible; iterate toward acceptable trade-offs.",
    }
    STUDY_NOTES[5] = {
        "Extracting Architectural Characteristics from Domain Concerns": "Translate business language (M&A, satisfaction) into measurable -ilities via tables and questions.",
        "Composite Architectural Characteristics": "Decompose “agility” or “reliability” into measurable pieces before designing.",
        "The Origin of Architecture Katas": "Practice architectures safely; katas = deliberate rehearsal like martial arts forms.",
        "Working with Katas": "Timeboxed, collaborative analysis; compare designs and missed trade-offs.",
        "Kata: Silicon Sandwiches": "Walk through explicit vs implicit characteristics, design vs architecture boundary.",
        "Explicit Characteristics": "Stated in requirements; still need prioritisation.",
        "Implicit Characteristics": "Assumed necessities—availability, security—must be surfaced.",
        "Design Versus Architecture and Trade-Offs": "Same behaviour might be pattern in code or structural style—cost and coupling decide.",
        "Limiting and Prioritizing Architectural Characteristics": "Worksheet + top-3 avoids Vasa-style overload.",
        "Case Study: The Vasa": "Too many conflicting goals capsized the ship; analogy for over-specified architectures.",
    }
    STUDY_NOTES[6] = {
        "Measuring Architecture Characteristics": "Define terms, decompose composites, build ubiquitous language.",
        "Operational Measures": "Use percentiles and models, not only averages; watch mobile UX budgets.",
        "The Many Flavors of Performance": "Latency, throughput, render budgets, K-weight—pick measures that match user journeys.",
        "Structural Measures": "No single “architecture quality” score; use focused metrics (e.g. CC) plus interpretation.",
        "Cyclomatic Complexity": "Counts decision paths; high values warrant refactor unless domain is inherently complex.",
        "Process Measures": "Test coverage, deploy frequency, lead time—tie “agility” to observable pipeline traits.",
        "Governance and Fitness Functions": "Automate checks so principles don’t erode under delivery pressure.",
        "Governing Architecture Characteristics": "Collaborate with devs on rules; avoid ivory-tower metrics.",
        "Fitness Functions": "Tests/monkeys/checklists encode architectural invariants into engineering workflow.",
    }
    STUDY_NOTES[7] = {
        "Architectural Quanta and Granularity": "Quantum = scoped bundle of deployment + cohesion + coupling rules; plural quanta.",
        "Domain-Driven Design's Bounded Context": "Local models per subdomain; avoid one giant shared Customer—reconcile at boundaries.",
        "Synchronous Communication": "Sync ties availability and latency across services; may reshape quantum boundaries.",
        "The Impact of Scoping": "Different subsystems can require different -ilities; one size rarely fits.",
        "Scoping and Architectural Style": "Monolith vs distributed choice flows from how many characteristic profiles you need.",
        "Kata: Going Green": "Example of multiple clusters of characteristics encouraging multiple quanta.",
        "Scoping and the Cloud": "Regions/zones/functions change how you attach operational characteristics to parts.",
    }
    STUDY_NOTES[8] = {
        "Defining Logical Components": "Rooms in a floor plan; leaf folders/namespaces often map to components.",
        "Logical Versus Physical Architecture": "Logical shows capabilities and flows; physical shows services/DB/UI—don’t skip logical thinking.",
        "Creating a Logical Architecture": "Iterative: identify, assign stories, analyse roles, check -ilities, refactor.",
        "Identifying Core Components": "Start with best guess; event/workflow/actor lenses; avoid perfectionism on v1.",
        "Assigning User Stories to Components": "Misfits reveal misplaced responsibilities or missing components.",
        "Analyzing Roles and Responsibilities": "Ensure each component’s job reads coherently; watch overlap.",
        "Analyzing Architectural Characteristics": "Split/merge components when -ilities demand different structures.",
        "Restructuring Components": "Expect to rename/split/merge as learning increases.",
        "Component Coupling": "Static and temporal links; less coupling ⇒ easier test and change.",
        "Static Coupling": "Call/import dependencies; measure Ca/Ce at component level.",
        "Temporal Coupling": "Order-sensitive workflows; harder to spot with tools.",
        "The Law of Demeter": "Limit knowledge of distant collaborators; push orchestration to the right owner.",
        "Case Study: Going, Going, Gone—Discovering Components": "Iterative component discovery for an auction system; multiple valid decompositions exist.",
    }

    for ch, start, end, title in CHAPTER_RANGES:
        build_chapter_doc(ch, start, end, title, lines)

    print("Done. Output:", OUT_DIR)


if __name__ == "__main__":
    main()
